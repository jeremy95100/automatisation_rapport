"""
Script pour ajouter le placeholder {daterap} dans les documents Word test.docx et test2.docx.
Ce placeholder permet à l'utilisateur de saisir la date du rapport.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_daterap_to_document(doc_path: Path):
    """Ajoute le placeholder {daterap} au début du document."""
    print(f"\nTraitement de {doc_path.name}...")

    if not doc_path.exists():
        print(f"   [ERREUR] Fichier introuvable: {doc_path}")
        return False

    # Créer une sauvegarde
    backup_path = doc_path.with_suffix('.docx.backup')
    if not backup_path.exists():
        import shutil
        shutil.copy2(doc_path, backup_path)
        print(f"   [OK] Sauvegarde creee: {backup_path.name}")

    # Charger le document
    doc = Document(str(doc_path))

    # Vérifier si {daterap} existe déjà
    text_content = "\n".join([p.text for p in doc.paragraphs])
    if "{daterap}" in text_content:
        print(f"   [INFO] Le placeholder {{daterap}} existe deja dans le document")
        return True

    # Ajouter le placeholder au début du document
    # Option 1: Ajouter dans un nouveau paragraphe en haut
    first_para = doc.paragraphs[0] if doc.paragraphs else None

    if first_para:
        # Insérer avant le premier paragraphe
        new_para = doc.add_paragraph()
        # Déplacer le nouveau paragraphe au début
        new_para._element.getparent().insert(0, new_para._element)

        # Ajouter le texte avec le placeholder
        run = new_para.add_run("Date du rapport: {daterap}")
        run.font.size = Pt(11)
        run.font.bold = True

        # Ajouter un saut de ligne
        doc.paragraphs[1].insert_paragraph_before("")

        print(f"   [OK] Placeholder {{daterap}} ajoute au debut du document")
    else:
        # Si le document est vide, créer un premier paragraphe
        para = doc.add_paragraph("Date du rapport: {daterap}")
        run = para.runs[0]
        run.font.size = Pt(11)
        run.font.bold = True
        print(f"   [OK] Placeholder {{daterap}} ajoute (document etait vide)")

    # Sauvegarder
    doc.save(str(doc_path))
    print(f"   [OK] Document sauvegarde avec succes")

    return True

def verify_placeholder(doc_path: Path):
    """Vérifie que le placeholder a bien été ajouté."""
    if not doc_path.exists():
        return False

    doc = Document(str(doc_path))
    text_content = "\n".join([p.text for p in doc.paragraphs])

    # Vérifier aussi dans les tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text_content += "\n" + p.text

    return "{daterap}" in text_content

def main():
    print("=" * 70)
    print("Ajout du placeholder {daterap} dans les templates")
    print("=" * 70)

    templates = [
        Path("test.docx"),
        Path("test2.docx")
    ]

    results = {}

    for template in templates:
        if template.exists():
            success = add_daterap_to_document(template)
            if success:
                verified = verify_placeholder(template)
                results[template.name] = verified
                if verified:
                    print(f"   [OK] Verification OK: {{daterap}} est bien present")
                else:
                    print(f"   [ATTENTION] Verification: {{daterap}} non detecte apres ajout")
        else:
            print(f"\n[ERREUR] {template.name} introuvable")
            results[template.name] = False

    print("\n" + "=" * 70)
    print("Resume:")
    for name, success in results.items():
        status = "[OK]" if success else "[ERREUR]"
        print(f"  {status} {name}")

    if all(results.values()):
        print("\n[OK] Tous les templates ont ete mis a jour avec succes!")
        print("\nLe placeholder {daterap} sera maintenant disponible pour:")
        print("  - L'interface web (via /placeholders)")
        print("  - Le mode interactif (via remplace_rapport.py)")
        print("  - L'API (via /generate)")
    else:
        print("\n[ATTENTION] Certains templates n'ont pas pu etre mis a jour")

    print("=" * 70)

if __name__ == "__main__":
    main()
