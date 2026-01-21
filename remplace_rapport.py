import argparse
import re
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Inches

PLACEHOLDER_PATTERN = re.compile(r"\{[^{}]+\}")
DEFAULT_ANALYSIS_TEMPLATE = ""
BACK_TOKEN = "__BACK__"
IMAGE_MARKER = re.compile(r"\[\[\s*IMG\s*:\s*([^\]]+?)\s*\]\]")


def remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def find_placeholders_in_order(doc):
    seen = set()
    ordered = []

    # Chercher dans les en-têtes de toutes les sections
    for section in doc.sections:
        header = section.header
        # Paragraphes dans l'en-tête
        for p in header.paragraphs:
            for ph in PLACEHOLDER_PATTERN.findall(p.text):
                if ph not in seen:
                    ordered.append(ph)
                    seen.add(ph)
        # Tableaux dans l'en-tête
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for ph in PLACEHOLDER_PATTERN.findall(p.text):
                            if ph not in seen:
                                ordered.append(ph)
                                seen.add(ph)

    # Chercher dans le corps du document
    for p in iter_all_paragraphs(doc):
        for ph in PLACEHOLDER_PATTERN.findall(p.text):
            if ph not in seen:
                ordered.append(ph)
                seen.add(ph)
    return ordered


def find_image_markers_in_order(doc):
    """Extract all [[IMG:key]] markers from the document in order."""
    seen = set()
    ordered = []

    # Search in headers of all sections
    for section in doc.sections:
        header = section.header
        # Paragraphs in header
        for p in header.paragraphs:
            for marker in IMAGE_MARKER.findall(p.text):
                if marker not in seen:
                    ordered.append(marker)
                    seen.add(marker)
        # Tables in header
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for marker in IMAGE_MARKER.findall(p.text):
                            if marker not in seen:
                                ordered.append(marker)
                                seen.add(marker)

    # Search in document body
    for p in iter_all_paragraphs(doc):
        for marker in IMAGE_MARKER.findall(p.text):
            if marker not in seen:
                ordered.append(marker)
                seen.add(marker)
    return ordered


def is_in_table_cell(paragraph):
    """Vérifie si le paragraphe est à l'intérieur d'une cellule de tableau."""
    parent = paragraph._element.getparent()
    while parent is not None:
        if parent.tag.endswith('}tc'):  # tc = table cell
            return True
        parent = parent.getparent()
    return False


def replace_in_runs(paragraph, mapping, allow_delete=True):
    """Remplace placeholders coupes en runs. Valeur vide -> supprime le paragraphe entier (sauf dans les cellules de tableau)."""
    if not mapping:
        return False
    original = paragraph.text
    new_text = original
    remove_entire = False

    for old, new in mapping.items():
        if old not in new_text:
            continue
        if new == "":
            # Ne supprimer le paragraphe que si on est HORS d'une cellule de tableau
            if allow_delete and not is_in_table_cell(paragraph):
                remove_entire = True
                break
            else:
                # Dans une cellule de tableau, on remplace juste par une chaîne vide
                new_text = new_text.replace(old, new)
        else:
            new_text = new_text.replace(old, new)

    if remove_entire:
        remove_paragraph(paragraph)
        return True
    if new_text == original:
        return False
    for run in list(paragraph.runs):
        run.text = ""
    paragraph.add_run(new_text)
    return False


def fill_with_mapping(text, mapping):
    out = text
    for old, new in mapping.items():
        out = out.replace(old, new)
    return out


def is_heading(paragraph):
    style_name = paragraph.style.name if paragraph.style else ""
    return style_name.startswith("Heading") or style_name.startswith("Titre")


def insert_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    new_para.add_run(text)
    return new_para


def remove_table(table):
    tbl_element = table._element
    parent = tbl_element.getparent()
    if parent is not None:
        parent.remove(tbl_element)
    if table in table._parent.tables:
        try:
            table._parent.tables.remove(table)
        except ValueError:
            pass


def remove_empty_paragraphs(doc):
    for p in list(iter_all_paragraphs(doc)):
        if not p.text or not p.text.strip():
            # Ne pas supprimer les paragraphes vides dans les cellules de tableau
            if not is_in_table_cell(p):
                remove_paragraph(p)


def remove_empty_sim_tables(doc, mapping):
    """Supprime les tableaux SIM dont tous les placeholders sont vides."""
    tables_to_remove = []

    print(f"[DEBUG remove_empty_sim_tables] Nombre de tableaux dans le document: {len(doc.tables)}")
    print(f"[DEBUG remove_empty_sim_tables] Clés SIM dans le mapping:")
    for i in range(1, 9):
        sim_keys = [f"{{operateur{i}}}", f"{{iccid{i}}}", f"{{imsi{i}}}", f"{{msisdn{i}}}", f"{{datesync{i}}}"]
        for key in sim_keys:
            if key in mapping:
                print(f"  {key} = '{mapping[key]}'")

    for table_idx, table in enumerate(doc.tables):
        # Vérifier si c'est un tableau SIM (a une ligne avec des placeholders operateur/iccid/imsi/msisdn/datesync)
        is_sim_table = False
        sim_index = None
        has_at_least_one_value = False

        for row in table.rows:
            row_text = "".join(cell.text for cell in row.cells)
            # Chercher des placeholders SIM indexés
            for i in range(1, 9):
                sim_keys = [f"{{operateur{i}}}", f"{{iccid{i}}}", f"{{imsi{i}}}", f"{{msisdn{i}}}", f"{{datesync{i}}}"]
                if any(key in row_text for key in sim_keys):
                    is_sim_table = True
                    sim_index = i
                    print(f"[DEBUG] Tableau {table_idx} identifié comme SIM {i}, row_text contient: {[k for k in sim_keys if k in row_text]}")
                    # Vérifier si AU MOINS UNE valeur est remplie dans le mapping pour cette carte SIM
                    for key in sim_keys:
                        value = mapping.get(key, "").strip()
                        print(f"[DEBUG]   Vérification {key} -> '{value}'")
                        if value:
                            has_at_least_one_value = True
                            print(f"[DEBUG]   => Valeur trouvée pour {key}!")
                            break
                    break
            if is_sim_table:
                break

        # Ne supprimer que si c'est un tableau SIM ET qu'aucune valeur n'est remplie
        if is_sim_table and not has_at_least_one_value:
            print(f"[DEBUG] Suppression du tableau SIM {sim_index} - aucune valeur remplie")
            tables_to_remove.append(table)
        elif is_sim_table:
            print(f"[DEBUG] Conservation du tableau SIM {sim_index} - au moins une valeur remplie")

    # Supprimer les tableaux marqués
    for table in tables_to_remove:
        remove_table(table)


def iter_block_items(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("tbl"):
            yield Table(child, doc)


def prompt_placeholders(placeholders):
    mapping = {}
    if not placeholders:
        print("Aucun placeholder { } trouve dans le document.")
        return mapping
    print("Placeholders trouves dans l'ordre :", ", ".join(placeholders))
    for ph in placeholders:
        value = input(f"Valeur pour {ph} (laisser vide = supprime le paragraphe contenant ce placeholder) : ")
        mapping[ph] = value
    return mapping 


def prompt_phrase_for_heading(title_text, mapping):
    print(f"Titre detecte : {title_text}")
    default_phrase = fill_with_mapping(DEFAULT_ANALYSIS_TEMPLATE, mapping)
    print(" 1) Saisir une phrase personnalisee")
    print(f" 2) {default_phrase}")
    print(" 3) Rien (supprime le titre et les tableaux qui suivent)")
    print(" b) Retour au titre precedent")
    choice = input("Choix (1/2/3/b, Enter=2) : ").strip().lower()
    if choice == "b":
        return BACK_TOKEN
    if choice == "1":
        return input("Votre phrase (laisser vide = suppression du titre/tableaux) : ").strip()
    if choice == "3":
        return ""
    if choice == "" or choice == "2":
        return default_phrase
    return default_phrase


def collect_headings_in_order(doc):
    return [p for p in iter_all_paragraphs(doc) if isinstance(p, Paragraph) and is_heading(p)]


def collect_heading_decisions(headings, mapping):
    decisions = []
    idx = 0
    while idx < len(headings):
        title_text = headings[idx].text.strip()
        res = prompt_phrase_for_heading(title_text, mapping)
        if res == BACK_TOKEN:
            if idx > 0:
                decisions.pop()
                idx -= 1
            else:
                print("Deja au premier titre, impossible de revenir en arriere.")
            continue
        if idx < len(decisions):
            decisions[idx] = res
        else:
            decisions.append(res)
        idx += 1
    return decisions


def apply_heading_decisions(doc, decisions):
    blocks = list(iter_block_items(doc))
    heading_idx = 0
    idx = 0
    while idx < len(blocks):
        blk = blocks[idx]
        if isinstance(blk, Paragraph) and is_heading(blk):
            decision = decisions[heading_idx]
            heading_idx += 1
            if decision == "__KEEP_TITLE_ONLY__":
                # Garder le titre mais ne rien ajouter en dessous
                pass
            elif decision:
                # Ajouter la phrase sous le titre
                insert_after(blk, decision, style=None)
            else:
                # Supprimer le titre et les tableaux qui suivent
                remove_paragraph(blk)
                blocks.pop(idx)
                j = idx
                while j < len(blocks):
                    nxt = blocks[j]
                    if isinstance(nxt, Paragraph) and is_heading(nxt):
                        break
                    if isinstance(nxt, Table):
                        remove_table(nxt)
                        blocks.pop(j)
                        continue
                    j += 1
                idx = j
                continue
        idx += 1


def default_heading_decisions(headings: List[Paragraph], mapping: Dict[str, str]) -> List[str]:
    """Genere une phrase par defaut sous chaque titre (non interactif)."""
    return [fill_with_mapping(DEFAULT_ANALYSIS_TEMPLATE, mapping) for _ in headings]

def insert_image_after(paragraph: Paragraph, image_path: str, width_inches: float = 3.0, text_before: str = "", text_after: str = ""):
    """Insere une image juste apres le paragraphe donne, avec optionnellement du texte avant/après."""
    # Vérifier que le fichier image existe
    img_path = Path(image_path)
    if not img_path.exists():
        print(f"ATTENTION: Image introuvable: {image_path}")
        return

    if text_before:
        new_p = insert_after(paragraph, text_before)
        paragraph = new_p
    new_p = insert_after(paragraph, "")
    run = new_p.add_run()
    try:
        run.add_picture(str(img_path), width=Inches(width_inches))
    except Exception as e:
        print(f"ERREUR lors de l'insertion de l'image {image_path}: {e}")
        return
    if text_after:
        insert_after(new_p, text_after)


def apply_images_after_headings(doc: Document, images_after: Dict[str, str], width_inches: float = 3.0,
                                per_image_widths: Optional[Dict[str, float]] = None,
                                image_texts: Optional[Dict[str, Dict[str, str]]] = None):
    """Insere des images apres les titres dont le texte matche exactement la cle du mapping."""
    if not images_after:
        return

    # Parcourir tous les paragraphes et identifier les headings
    all_paras = list(doc.paragraphs)
    for idx, para in enumerate(all_paras):
        if is_heading(para):
            heading_text = para.text.strip()
            img = images_after.get(heading_text)
            if img:
                w = per_image_widths.get(heading_text) if per_image_widths else None
                text_data = image_texts.get(heading_text, {}) if image_texts else {}
                text_before = text_data.get("before", "") if text_data.get("position") == "before" else ""
                text_after = text_data.get("after", "") if text_data.get("position") == "after" else ""

                # Trouver le paragraphe suivant qui n'est pas un heading
                target_para = para
                if idx + 1 < len(all_paras):
                    next_para = all_paras[idx + 1]
                    if not is_heading(next_para):
                        # Le prochain paragraphe est probablement la phrase automatique
                        target_para = next_para

                insert_image_after(target_para, img, w or width_inches, text_before, text_after)


def apply_images_after_paragraphs(doc: Document, images_after: Dict[str, str], width_inches: float = 3.0,
                                  per_image_widths: Optional[Dict[str, float]] = None,
                                  image_texts: Optional[Dict[str, Dict[str, str]]] = None):
    """Insere des images apres les paragraphes dont le texte matche exactement la cle du mapping."""
    if not images_after:
        return

    # Utiliser un set pour éviter les insertions multiples du même paragraphe
    processed = set()
    for p in iter_all_paragraphs(doc):
        # Créer un identifiant unique basé sur le contenu et la position
        para_id = id(p._element)
        if para_id in processed:
            continue

        key = p.text.strip()
        img = images_after.get(key)
        if img:
            w = per_image_widths.get(key) if per_image_widths else None
            text_data = image_texts.get(key, {}) if image_texts else {}
            text_before = text_data.get("before", "") if text_data.get("position") == "before" else ""
            text_after = text_data.get("after", "") if text_data.get("position") == "after" else ""
            insert_image_after(p, img, w or width_inches, text_before, text_after)
            processed.add(para_id) 


def apply_images_at_markers(doc: Document, images_at_markers: Dict[str, str], width_inches: float = 3.0,
                            per_image_widths: Optional[Dict[str, float]] = None,
                            image_texts: Optional[Dict[str, Dict[str, str]]] = None):
    """Remplace les marqueurs [[IMG:cle]] par l'image correspondante inseree a cet endroit."""
    if not images_at_markers:
        return
    for p in iter_all_paragraphs(doc):
        full_text = "".join(run.text for run in p.runs)
        if "[[IMG" not in full_text:
            continue
        parts = IMAGE_MARKER.split(full_text)  # [texte, cle, texte, cle, ...]
        for run in list(p.runs):
            run.text = ""
        for idx, chunk in enumerate(parts):
            if idx % 2 == 0:
                if chunk:
                    p.add_run(chunk)
            else:
                key = chunk.strip()
                img_path = images_at_markers.get(key)
                if img_path:
                    # Vérifier que le fichier existe
                    img_file = Path(img_path)
                    if not img_file.exists():
                        print(f"ATTENTION: Image introuvable pour le marqueur '{key}': {img_path}")
                        p.add_run(f"[[IMG:{key} - INTROUVABLE]]")
                        continue

                    text_data = image_texts.get(key, {}) if image_texts else {}
                    # Add text before image
                    if text_data.get("position") == "before" and text_data.get("before"):
                        p.add_run(text_data["before"] + " ")
                    # Add image
                    try:
                        w = per_image_widths.get(key) if per_image_widths else None
                        p.add_run().add_picture(str(img_file), width=Inches(w or width_inches))
                    except Exception as e:
                        print(f"ERREUR lors de l'insertion de l'image pour le marqueur '{key}': {e}")
                        p.add_run(f"[[IMG:{key} - ERREUR]]")
                        continue
                    # Add text after image
                    if text_data.get("position") == "after" and text_data.get("after"):
                        p.add_run(" " + text_data["after"])
                else:
                    p.add_run(f"[[IMG:{key}]]")


def apply_heading_content_blocks(doc: Document, heading_content: Dict[str, List[Dict]], default_width_inches: float = 3.0):
    """Insere les blocs de contenu (texte/images) apres chaque heading specifie."""
    if not heading_content:
        return

    # Parcourir tous les paragraphes pour trouver les headings
    all_paras = list(doc.paragraphs)
    for idx, para in enumerate(all_paras):
        if is_heading(para):
            heading_text = para.text.strip()
            blocks = heading_content.get(heading_text, [])
            if not blocks:
                continue

            # Trouver le paragraphe apres le heading (probablement la phrase automatique)
            target_para = para
            if idx + 1 < len(all_paras):
                next_para = all_paras[idx + 1]
                if not is_heading(next_para):
                    target_para = next_para

            # Inserer chaque bloc dans l'ordre
            current_para = target_para
            for block in blocks:
                if block.get("type") == "text":
                    content = block.get("content", "")
                    if content:
                        new_para = insert_after(current_para, content)
                        current_para = new_para
                elif block.get("type") == "image":
                    src = block.get("src", "")
                    if src:
                        width = block.get("width") or default_width_inches
                        img_path = Path(src)
                        if not img_path.exists():
                            print(f"ATTENTION: Image introuvable: {src}")
                            continue

                        # Creer un nouveau paragraphe pour l'image
                        new_para = insert_after(current_para, "")
                        run = new_para.add_run()
                        try:
                            run.add_picture(str(img_path), width=Inches(width))
                            current_para = new_para
                        except Exception as e:
                            print(f"ERREUR lors de l'insertion de l'image {src}: {e}")


def process_document(input_path, output_path, mapping_override: Optional[Dict[str, str]] = None,
                     decisions_override: Optional[List[str]] = None, interactive: bool = True,
                     heading_content: Optional[Dict[str, List[Dict]]] = None,
                     images_at_markers: Optional[Dict[str, str]] = None,
                     image_width_inches: float = 3.0,
                     images_at_markers_sizes: Optional[Dict[str, float]] = None):
    input_path = Path(input_path)
    output_path = Path(output_path)

    doc = Document(str(input_path))

    placeholders = find_placeholders_in_order(doc)
    if mapping_override is not None:
        mapping = mapping_override
    elif interactive:
        mapping = prompt_placeholders(placeholders)
    else:
        mapping = {}

    # Remplacer dans les en-têtes
    for section in doc.sections:
        header = section.header
        # Paragraphes dans l'en-tête
        for p in header.paragraphs:
            replace_in_runs(p, mapping)
        # Tableaux dans l'en-tête
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_runs(p, mapping)

    # Supprimer les tableaux SIM vides AVANT le remplacement des placeholders
    remove_empty_sim_tables(doc, mapping)

    # Remplacer dans le corps du document
    for p in list(iter_all_paragraphs(doc)):
        removed = replace_in_runs(p, mapping)
        if removed:
            continue
    remove_empty_paragraphs(doc)

    headings = collect_headings_in_order(doc)
    if decisions_override is not None:
        decisions = decisions_override
    elif interactive:
        decisions = collect_heading_decisions(headings, mapping)
    else:
        decisions = default_heading_decisions(headings, mapping)

    apply_heading_decisions(doc, decisions)
    remove_empty_paragraphs(doc)

    # Insertion des blocs de contenu (texte + images) après les headings
    if heading_content:
        apply_heading_content_blocks(doc, heading_content, default_width_inches=image_width_inches)

    # Insertion d'images sur les marqueurs
    if images_at_markers:
        apply_images_at_markers(doc, images_at_markers, width_inches=image_width_inches,
                                per_image_widths=images_at_markers_sizes)

    doc.save(str(output_path))
    print(f"Document genere : {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Automation interactive pour Word (placeholders + phrases sous titres)")
    parser.add_argument("--input", default="test.docx", help="Fichier Word source")
    parser.add_argument("--output", default="test_sortie.docx", help="Fichier Word de sortie")
    args = parser.parse_args()
    process_document(args.input, args.output)


if __name__ == "__main__":
    main()
