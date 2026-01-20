import asyncio
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import io

import mammoth
from docx import Document
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, RedirectResponse, StreamingResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from PIL import Image
import pillow_heif

from remplace_rapport import (
    collect_headings_in_order,
    DEFAULT_ANALYSIS_TEMPLATE,
    default_heading_decisions,
    find_placeholders_in_order,
    find_image_markers_in_order,
    process_document,
)

TEMPLATES = {
    "test": Path("test.docx"),
    "test2": Path("test2.docx"),
    "test3": Path("test3.docx"),
}
OUTPUT_PATHS = {name: path.with_name(f"{path.stem}_sortie.docx") for name, path in TEMPLATES.items()}
FRONTEND_DIR = Path("frontend")
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

app = FastAPI(title="Rapport auto - API")  # HEIC support enabled
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

listeners: List[asyncio.Queue] = []
doc_lock = asyncio.Lock()


class ImageTextData(BaseModel):
    before: str = ""
    after: str = ""
    position: str = "none"  # "before", "after", or "none"

class ContentBlock(BaseModel):
    type: str  # "text" or "image"
    content: Optional[str] = None  # for text blocks
    src: Optional[str] = None  # for image blocks
    width: Optional[float] = None  # for image blocks (in inches)

class GeneratePayload(BaseModel):
    template: Optional[str] = None
    overwrite: bool = False
    mapping: Dict[str, str] = {}
    decisions: Optional[List[str]] = None
    heading_content: Dict[str, List[ContentBlock]] = {}  # nouvelle structure
    images_at_markers: Dict[str, str] = {}
    image_width_inches: Optional[float] = None
    images_at_markers_sizes: Dict[str, float] = {}


def available_templates() -> List[str]:
    # On liste toutes les trames déclarées ; l'existence est vérifiée plus tard
    return list(TEMPLATES.keys())


def get_template_paths(selected: Optional[str] = None) -> Tuple[Path, Path, str]:
    if selected:
        if selected not in TEMPLATES:
            raise HTTPException(status_code=404, detail=f"Trame inconnue: {selected}")
        src = TEMPLATES[selected]
        if not src.exists():
            raise HTTPException(status_code=404, detail=f"Source file not found: {src}")
        return src, OUTPUT_PATHS[selected], selected
    for name, src in TEMPLATES.items():
        if src.exists():
            return src, OUTPUT_PATHS[name], name
    raise HTTPException(status_code=404, detail="Aucune trame disponible")


def convert_to_html(docx_path: Path) -> str:
    if not docx_path.exists():
        raise HTTPException(status_code=404, detail="Document introuvable pour l'aperçu HTML")

    # Extraire les en-têtes avec python-docx
    doc = Document(str(docx_path))
    headers_html = ""

    # Récupérer les en-têtes de toutes les sections
    for section in doc.sections:
        header = section.header

        # Vérifier les paragraphes dans l'en-tête
        for para in header.paragraphs:
            if para.text.strip():
                style = 'style="text-align: center; font-weight: bold; margin-bottom: 10px;"'
                headers_html += f'<div {style}>{para.text}</div>\n'

        # Vérifier les tableaux dans l'en-tête (cas de test.docx et test2.docx)
        for table in header.tables:
            headers_html += '<table style="width: 100%; border-collapse: collapse; margin-bottom: 10px;">\n'
            for row in table.rows:
                headers_html += '  <tr>\n'
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        # Déterminer l'alignement et le style selon le contenu
                        if "Rapport" in cell_text or "Art." in cell_text:
                            style = 'style="text-align: center; font-weight: bold; padding: 5px; vertical-align: middle;"'
                        else:
                            style = 'style="text-align: right; padding: 5px; vertical-align: middle;"'
                        # Préserver les sauts de ligne
                        cell_html = cell_text.replace('\n', '<br>')
                        headers_html += f'    <td {style}>{cell_html}</td>\n'
                    else:
                        headers_html += '    <td style="padding: 5px;"></td>\n'
                headers_html += '  </tr>\n'
            headers_html += '</table>\n'

    # Si pas d'en-têtes trouvés dans les sections, chercher dans les premiers paragraphes
    if not headers_html:
        for para in doc.paragraphs[:5]:
            text = para.text.strip()
            if text and "{" in text:
                style = 'style="text-align: center; font-weight: bold; margin-bottom: 10px;"'
                headers_html += f'<div {style}>{text}</div>\n'

    # Convertir le reste du document avec mammoth
    with open(docx_path, "rb") as f:
        result = mammoth.convert_to_html(f)

    # Combiner l'en-tête et le contenu
    if headers_html:
        full_html = f'<div style="border-bottom: 2px solid #333; padding-bottom: 20px; margin-bottom: 20px;">\n{headers_html}</div>\n{result.value}'
    else:
        full_html = result.value

    return full_html


async def broadcast(message: str) -> None:
    for q in list(listeners):
        await q.put(message)


async def event_stream():
    queue: asyncio.Queue = asyncio.Queue()
    listeners.append(queue)
    try:
        while True:
            msg = await queue.get()
            yield f"data: {msg}\n\n"
    finally:
        listeners.remove(queue)


@app.get("/")
def root():
    if FRONTEND_DIR.exists():
        return RedirectResponse(url="/ui/")
    return {"message": "API running", "routes": ["/templates", "/placeholders", "/generate", "/preview", "/events"]}


@app.get("/templates")
def list_templates():
    templates = available_templates()
    default = None
    for name in templates:
        if TEMPLATES[name].exists():
            default = name
            break
    if default is None:
        default = templates[0] if templates else None
    return {"templates": templates, "default": default}


@app.get("/placeholders")
def get_placeholders(template: Optional[str] = None):
    src_path, _, template_key = get_template_paths(template)
    doc = Document(str(src_path))
    placeholders = find_placeholders_in_order(doc)
    headings = [p.text for p in collect_headings_in_order(doc)]
    markers = find_image_markers_in_order(doc)
    return {
        "template": template_key,
        "templates": available_templates(),
        "placeholders": placeholders,
        "headings": headings,
        "markers": markers,
        "default_template": DEFAULT_ANALYSIS_TEMPLATE,
    }


@app.post("/generate")
async def generate(payload: GeneratePayload):
    src_path, output_path, template_key = get_template_paths(payload.template)
    mapping = payload.mapping or {}
    if payload.overwrite:
        if output_path.exists():
            output_path.unlink()
    doc = Document(str(src_path))
    placeholders_in_doc = find_placeholders_in_order(doc)
    for missing in placeholders_in_doc:
        mapping.setdefault(missing, "")
    headings = collect_headings_in_order(doc)
    decisions = payload.decisions
    if decisions is None:
        decisions = default_heading_decisions(headings, mapping)
    else:
        # "__DEFAULT__" => phrase auto, "" => rien, autre => texte personnalisé
        resolved = []
        defaults = default_heading_decisions(headings, mapping)
        for idx, h in enumerate(headings):
            choice = payload.decisions[idx] if idx < len(payload.decisions) else "__DEFAULT__"
            if choice == "__DEFAULT__" or choice is None:
                resolved.append(defaults[idx])
            else:
                resolved.append(choice)
        decisions = resolved

    def resolve_path(path_str: str) -> str:
        if path_str.startswith("/uploads/"):
            return str(UPLOAD_DIR / Path(path_str).name)
        return path_str

    # Convert heading_content blocks to resolved paths
    heading_content_resolved = {}
    for heading, blocks in (payload.heading_content or {}).items():
        heading_content_resolved[heading] = []
        for block in blocks:
            block_dict = block.model_dump()
            if block.type == "image" and block.src:
                block_dict["src"] = resolve_path(block.src)
            heading_content_resolved[heading].append(block_dict)

    # Resolve marker images
    markers_resolved = {}
    for k, v in (payload.images_at_markers or {}).items():
        markers_resolved[k] = resolve_path(v)

    # Debug logging
    print(f"[DEBUG] mapping: {mapping}")
    print(f"[DEBUG] decisions: {decisions}")
    print(f"[DEBUG] heading_content_resolved: {heading_content_resolved}")
    print(f"[DEBUG] markers_resolved: {markers_resolved}")

    async with doc_lock:
        process_document(
            src_path,
            output_path,
            mapping_override=mapping,
            decisions_override=decisions,
            interactive=False,
            heading_content=heading_content_resolved,
            images_at_markers=markers_resolved,
            image_width_inches=payload.image_width_inches or 3.0,
            images_at_markers_sizes=payload.images_at_markers_sizes or {},
        )
    await broadcast("updated")
    return {"status": "ok", "template": template_key, "output": str(output_path), "pdf": None}


@app.get("/download")
async def download(template: Optional[str] = None):
    """Télécharge le fichier Word généré"""
    _, output_path, _ = get_template_paths(template)
    if not output_path.exists():
        raise HTTPException(status_code=404, detail="Fichier de sortie non trouvé. Générez d'abord le document.")
    return FileResponse(
        path=str(output_path),
        filename=output_path.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.get("/preview")
async def preview(template: Optional[str] = None):
    src_path, output_path, template_key = get_template_paths(template)
    target = output_path if output_path.exists() else src_path
    if not target.exists():
        raise HTTPException(status_code=404, detail="Aucun fichier de reference disponible")
    async with doc_lock:
        html = await asyncio.get_running_loop().run_in_executor(None, convert_to_html, target)
    return HTMLResponse(html)


@app.get("/preview/html")
async def preview_html(template: Optional[str] = None):
    src_path, output_path, _ = get_template_paths(template)
    target = output_path if output_path.exists() else src_path
    if not target.exists():
        raise HTTPException(status_code=404, detail="Aucun fichier de reference disponible")
    async with doc_lock:
        html = await asyncio.get_running_loop().run_in_executor(None, convert_to_html, target)
    return HTMLResponse(html)


@app.get("/events")
async def events():
    return StreamingResponse(event_stream(), media_type="text/event-stream")


@app.post("/upload")
async def upload_image(file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="Nom de fichier manquant")

    contents = await file.read()
    original_filename = file.filename
    filename_lower = original_filename.lower()

    print(f"[UPLOAD DEBUG] Fichier reçu: {original_filename}, extension détectée: {filename_lower}")

    # Convertir HEIC en JPEG si nécessaire
    if filename_lower.endswith('.heic') or filename_lower.endswith('.heif'):
        try:
            # Enregistrer le support HEIF
            pillow_heif.register_heif_opener()

            # Ouvrir l'image HEIC
            heic_image = Image.open(io.BytesIO(contents))

            # Convertir en RGB si nécessaire (HEIC peut être en RGBA)
            if heic_image.mode in ('RGBA', 'LA', 'P'):
                rgb_image = Image.new('RGB', heic_image.size, (255, 255, 255))
                if heic_image.mode == 'P':
                    heic_image = heic_image.convert('RGBA')
                rgb_image.paste(heic_image, mask=heic_image.split()[-1] if heic_image.mode == 'RGBA' else None)
                heic_image = rgb_image
            elif heic_image.mode != 'RGB':
                heic_image = heic_image.convert('RGB')

            # Nouveau nom de fichier en .jpg (garder le nom original sans l'extension)
            new_filename = Path(original_filename).stem + '.jpg'
            
            dest = UPLOAD_DIR / new_filename

            # Sauvegarder en JPEG avec bonne qualité
            heic_image.save(dest, 'JPEG', quality=95, optimize=True)

            web_path = f"/uploads/{new_filename}"
            print(f"[UPLOAD] HEIC converti: {original_filename} -> {new_filename}")
            return JSONResponse({"path": web_path})
        except Exception as e:
            print(f"[UPLOAD ERROR] Erreur conversion HEIC {original_filename}: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Erreur conversion HEIC: {str(e)}")
    else:
        # Pour les autres formats, sauvegarder directement
        dest = UPLOAD_DIR / file.filename
        dest.write_bytes(contents)
        web_path = f"/uploads/{file.filename}"
        print(f"[UPLOAD] Image sauvegardée: {file.filename}")
        return JSONResponse({"path": web_path})


if FRONTEND_DIR.exists():
    app.mount("/ui", StaticFiles(directory=str(FRONTEND_DIR), html=True), name="ui")
app.mount("/uploads", StaticFiles(directory=str(UPLOAD_DIR), html=False), name="uploads")

